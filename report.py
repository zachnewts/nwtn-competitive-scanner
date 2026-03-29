"""
Generates a polished, executive-quality Word doc from competitive signals.

Uses python-docx to create a corporate briefing document with:
- NWTN.AI wordmark as styled text (Lora serif, ember accent period)
- Title page with confidentiality notice
- Running header with wordmark + confidential marking on content pages
- Navy accents, Calibri body, structured sections, color-coded badges
- Claude-generated executive summary
"""

import os
import logging
from datetime import datetime

import anthropic
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from models import CompetitiveSignal, Classification, ScanReport
import config

logger = logging.getLogger(__name__)

# ─── Design Tokens ───────────────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x36, 0x5D)
EMBER = RGBColor(0xC4, 0x5D, 0x3E)     # NWTN brand accent
TEXT_DARK = RGBColor(0x4A, 0x4A, 0x4A)  # Wordmark text on light bg
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xEE, 0xEE, 0xEE)
CONFIDENTIAL_GRAY = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

BADGE_COLORS = {
    Classification.DIRECT_COMPETITOR: RGBColor(0xC0, 0x39, 0x2B),
    Classification.ADJACENT_THREAT: RGBColor(0xE6, 0x7E, 0x22),
    Classification.POTENTIAL_PARTNER: RGBColor(0x27, 0xAE, 0x60),
    Classification.IRRELEVANT: RGBColor(0x99, 0x99, 0x99),
}

BADGE_LABELS = {
    Classification.DIRECT_COMPETITOR: "DIRECT COMPETITOR",
    Classification.ADJACENT_THREAT: "ADJACENT THREAT",
    Classification.POTENTIAL_PARTNER: "POTENTIAL PARTNER",
    Classification.IRRELEVANT: "IRRELEVANT",
}

FONT_BODY = "Calibri"
FONT_WORDMARK = "Lora"  # Serif for the wordmark; falls back if not installed


def generate_report(
    signals: list[CompetitiveSignal],
    new_count: int,
    updated_count: int,
    queries_run: int,
    total_raw_results: int,
) -> ScanReport:
    """Generate a polished Word doc report from scored competitive signals."""
    os.makedirs(config.REPORT_DIR, exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M")
    filename = f"competitive_scan_{timestamp}.docx"
    file_path = os.path.join(config.REPORT_DIR, filename)

    print(f"[REPORT] Generating report: {filename}")

    new_signals = [s for s in signals if s.is_new]
    updated_signals = [s for s in signals if not s.is_new]

    exec_summary = _generate_executive_summary(signals, new_count, updated_count)

    doc = Document()
    _set_default_font(doc)

    # ── Title Page (first section) ────────────────────────────────────────
    _set_margins(doc.sections[0])
    _add_title_page(doc, exec_summary)

    # ── Content Pages (new section with running header) ───────────────────
    doc.add_section()
    content_section = doc.sections[-1]
    _set_margins(content_section)
    _add_running_header(content_section)

    if new_signals:
        _add_section_heading(doc, "New Signals")
        for i, signal in enumerate(new_signals):
            _add_signal_entry(doc, signal)
            if i < len(new_signals) - 1:
                _add_divider(doc)

    if updated_signals:
        _add_spacer(doc)
        _add_section_heading(doc, "Updated Signals")
        _add_muted_text(doc, "Previously identified entities with new activity this scan.")
        for i, signal in enumerate(updated_signals):
            _add_signal_entry_light(doc, signal)
            if i < len(updated_signals) - 1:
                _add_divider(doc)

    _add_spacer(doc)
    _add_metadata_section(doc, queries_run, total_raw_results, new_count, updated_count)
    _add_footer(doc)

    doc.save(file_path)
    print(f"[REPORT] Saved: {file_path}")

    high_overlap = sum(1 for s in signals if s.overlap_score >= config.HIGH_PRIORITY_THRESHOLD)
    slack_summary = (
        f"Competitive scan complete: {new_count} new signals, "
        f"{updated_count} updated. {high_overlap} high-overlap."
    )

    return ScanReport(
        queries_run=queries_run,
        total_results_found=total_raw_results,
        new_signals=new_count,
        updated_signals=updated_count,
        summary=slack_summary,
        file_path=file_path,
    )


# ═════════════════════════════════════════════════════════════════════════════
#  EXECUTIVE SUMMARY (Claude-generated)
# ═════════════════════════════════════════════════════════════════════════════

def _generate_executive_summary(
    signals: list[CompetitiveSignal],
    new_count: int,
    updated_count: int,
) -> str:
    """Ask Claude to write a 2-3 sentence executive summary."""
    if not config.ANTHROPIC_API_KEY:
        return f"{new_count} new competitive signals identified. Review details below."

    signal_summaries = []
    for s in signals:
        signal_summaries.append(
            f"- {s.company_name} ({s.classification.value}, composite: {s.overlap_score}): "
            f"{s.description}"
        )

    prompt = (
        f"You are writing the executive summary for NWTN AI's weekly competitive "
        f"landscape scan. Write exactly 2-3 sentences summarizing what matters most.\n\n"
        f"Signals found this scan ({new_count} new, {updated_count} updated):\n"
        + "\n".join(signal_summaries)
        + "\n\nNWTN AI's positioning: " + config.NWTN_POSITIONING.strip()
        + "\n\nWrite the summary. Be direct and strategic — what should NWTN's "
        f"founder pay attention to? No preamble, just the summary sentences."
    )

    try:
        client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=200,
            messages=[{"role": "user", "content": prompt}],
        )
        print("  → Executive summary generated")
        return response.content[0].text.strip()
    except Exception as e:
        logger.error(f"Executive summary generation failed: {e}")
        return f"{new_count} new competitive signals identified. Review details below."


# ═════════════════════════════════════════════════════════════════════════════
#  WORDMARK
# ═════════════════════════════════════════════════════════════════════════════

def _add_wordmark(para, size: Pt, text_color: RGBColor = TEXT_DARK) -> None:
    """Add the NWTN.AI wordmark as styled text runs on an existing paragraph.

    Wordmark spec: "NWTN" in bold + "." in ember #C45D3E + "AI" same as NWTN.
    Font: Lora (serif), weight 700, tight letter-spacing.
    """
    # "NWTN"
    nwtn = para.add_run("NWTN")
    nwtn.font.name = FONT_WORDMARK
    nwtn.font.size = size
    nwtn.font.color.rgb = text_color
    nwtn.bold = True
    _add_character_spacing(nwtn, -10)

    # "." — ember accent
    dot = para.add_run(".")
    dot.font.name = FONT_WORDMARK
    dot.font.size = size
    dot.font.color.rgb = EMBER
    dot.bold = True
    _add_character_spacing(dot, -10)

    # "AI"
    ai = para.add_run("AI")
    ai.font.name = FONT_WORDMARK
    ai.font.size = size
    ai.font.color.rgb = text_color
    ai.bold = True
    _add_character_spacing(ai, -10)


# ═════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════════

def _add_title_page(doc: Document, exec_summary: str) -> None:
    """Build the title/cover page with wordmark, title, and confidentiality."""
    # Vertical space to push content down
    for _ in range(4):
        _add_spacer(doc)

    # Wordmark — large, centered
    wordmark_para = doc.add_paragraph()
    wordmark_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    wordmark_para.paragraph_format.space_after = Pt(20)
    _add_wordmark(wordmark_para, Pt(36))

    # Report title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Competitive Landscape Scan")
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY
    run.font.name = FONT_BODY
    run.bold = True

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_para.paragraph_format.space_after = Pt(24)
    run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(11)
    run.font.color.rgb = MID_GRAY
    run.font.name = FONT_BODY

    # Navy divider
    _add_colored_divider(doc, NAVY)

    # Executive Summary label
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(14)
    label.paragraph_format.space_after = Pt(6)
    run = label.add_run("EXECUTIVE SUMMARY")
    run.font.size = Pt(8)
    run.font.color.rgb = NAVY
    run.font.name = FONT_BODY
    run.bold = True
    _add_character_spacing(run, 80)

    # Executive Summary body
    summary_para = doc.add_paragraph()
    summary_para.paragraph_format.space_after = Pt(12)
    run = summary_para.add_run(exec_summary)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    run.font.name = FONT_BODY
    run.italic = True

    # Push confidentiality notice toward bottom
    for _ in range(6):
        _add_spacer(doc)

    # Confidentiality notice
    _add_colored_divider(doc, LIGHT_GRAY)
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    conf_para.paragraph_format.space_before = Pt(8)
    conf_para.paragraph_format.space_after = Pt(2)
    run = conf_para.add_run("CONFIDENTIAL")
    run.font.size = Pt(7)
    run.font.color.rgb = CONFIDENTIAL_GRAY
    run.font.name = FONT_BODY
    run.bold = True
    _add_character_spacing(run, 100)

    conf_body = doc.add_paragraph()
    conf_body.alignment = WD_ALIGN_PARAGRAPH.LEFT
    conf_body.paragraph_format.space_after = Pt(4)
    run = conf_body.add_run(
        "This document is prepared exclusively for NWTN AI internal use. "
        "It contains proprietary competitive intelligence and should not be "
        "distributed, reproduced, or shared with external parties without "
        "written authorization."
    )
    run.font.size = Pt(7.5)
    run.font.color.rgb = CONFIDENTIAL_GRAY
    run.font.name = FONT_BODY

    # Page break after title page
    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════════
#  RUNNING HEADER (content pages)
# ═════════════════════════════════════════════════════════════════════════════

def _add_running_header(section) -> None:
    """Add a running header with wordmark + confidential notice to content pages."""
    header = section.header
    header.is_linked_to_previous = False

    # Single paragraph with wordmark left, confidential right (via tab stops)
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    para.paragraph_format.space_after = Pt(2)

    # Wordmark (small)
    _add_wordmark(para, Pt(9), TEXT_DARK)

    # Tab + right-aligned confidential text
    tab_run = para.add_run("\t\t")
    tab_run.font.size = Pt(7)

    conf_run = para.add_run("CONFIDENTIAL — NWTN AI INTERNAL USE ONLY")
    conf_run.font.size = Pt(6.5)
    conf_run.font.color.rgb = CONFIDENTIAL_GRAY
    conf_run.font.name = FONT_BODY
    _add_character_spacing(conf_run, 40)

    # Set right-aligned tab stop
    pPr = para._p.get_or_add_pPr()
    tabs = pPr.makeelement(qn("w:tabs"), {})
    tab = tabs.makeelement(
        qn("w:tab"),
        {qn("w:val"): "right", qn("w:pos"): "9360"},  # ~6.5" at narrow margins
    )
    tabs.append(tab)
    pPr.append(tabs)

    # Thin navy line under header
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "4",
            qn("w:color"): f"{NAVY[0]:02X}{NAVY[1]:02X}{NAVY[2]:02X}",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


# ═════════════════════════════════════════════════════════════════════════════
#  DOCUMENT STRUCTURE
# ═════════════════════════════════════════════════════════════════════════════

def _set_default_font(doc: Document) -> None:
    """Set Calibri as the default font for the entire document."""
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_GRAY
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)


def _set_margins(section) -> None:
    """Set professional narrow margins on a section."""
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def _add_section_heading(doc: Document, title: str) -> None:
    """Add a section heading with navy accent."""
    _add_colored_divider(doc, NAVY)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(10)
    run = heading.add_run(title.upper())
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    run.font.name = FONT_BODY
    run.bold = True
    _add_character_spacing(run, 60)


def _add_signal_entry(doc: Document, signal: CompetitiveSignal) -> None:
    """Add a full signal entry with all details."""
    # Company name + classification badge + score on same line
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_before = Pt(8)
    name_para.paragraph_format.space_after = Pt(4)

    name_run = name_para.add_run(signal.company_name)
    name_run.font.size = Pt(13)
    name_run.font.color.rgb = NAVY
    name_run.font.name = FONT_BODY
    name_run.bold = True

    name_para.add_run("   ")

    badge_run = name_para.add_run(BADGE_LABELS.get(signal.classification, "UNKNOWN"))
    badge_run.font.size = Pt(7.5)
    badge_run.font.color.rgb = BADGE_COLORS.get(signal.classification, MID_GRAY)
    badge_run.font.name = FONT_BODY
    badge_run.bold = True
    _add_character_spacing(badge_run, 40)

    name_para.add_run("   ")

    score_run = name_para.add_run(f"{signal.overlap_score:.0%} overlap")
    score_run.font.size = Pt(9)
    score_run.font.color.rgb = MID_GRAY
    score_run.font.name = FONT_BODY

    # Description
    _add_body_text(doc, signal.description)

    # Positioning
    if signal.positioning and signal.positioning != "N/A":
        pos_para = doc.add_paragraph()
        pos_para.paragraph_format.space_after = Pt(2)
        label_run = pos_para.add_run("Positioning: ")
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = NAVY
        label_run.font.name = FONT_BODY
        label_run.bold = True
        val_run = pos_para.add_run(signal.positioning)
        val_run.font.size = Pt(9)
        val_run.font.color.rgb = DARK_GRAY
        val_run.font.name = FONT_BODY
        val_run.italic = True

    # Scoring breakdown
    scores_para = doc.add_paragraph()
    scores_para.paragraph_format.space_after = Pt(2)
    label_run = scores_para.add_run("Scores: ")
    label_run.font.size = Pt(8)
    label_run.font.color.rgb = NAVY
    label_run.font.name = FONT_BODY
    label_run.bold = True
    scores_text = (
        f"Market {signal.market_overlap:.0%}  ·  "
        f"Service {signal.service_overlap:.0%}  ·  "
        f"Positioning {signal.positioning_overlap:.0%}  ·  "
        f"Credibility {signal.credibility_score:.0%}"
    )
    scores_run = scores_para.add_run(scores_text)
    scores_run.font.size = Pt(8)
    scores_run.font.color.rgb = MID_GRAY
    scores_run.font.name = FONT_BODY

    # Why it matters
    reasoning_para = doc.add_paragraph()
    reasoning_para.paragraph_format.space_after = Pt(2)
    label_run = reasoning_para.add_run("Why it matters: ")
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = NAVY
    label_run.font.name = FONT_BODY
    label_run.bold = True
    val_run = reasoning_para.add_run(signal.overlap_reasoning)
    val_run.font.size = Pt(9)
    val_run.font.color.rgb = DARK_GRAY
    val_run.font.name = FONT_BODY

    # Source URL
    url_para = doc.add_paragraph()
    url_para.paragraph_format.space_after = Pt(6)
    run = url_para.add_run(signal.url)
    run.font.size = Pt(7.5)
    run.font.color.rgb = MID_GRAY
    run.font.name = FONT_BODY


def _add_signal_entry_light(doc: Document, signal: CompetitiveSignal) -> None:
    """Add a lighter signal entry for updated/re-found signals."""
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_before = Pt(6)
    name_para.paragraph_format.space_after = Pt(3)

    name_run = name_para.add_run(signal.company_name)
    name_run.font.size = Pt(11)
    name_run.font.color.rgb = DARK_GRAY
    name_run.font.name = FONT_BODY
    name_run.bold = True

    name_para.add_run("   ")

    badge_run = name_para.add_run(BADGE_LABELS.get(signal.classification, "UNKNOWN"))
    badge_run.font.size = Pt(7)
    badge_run.font.color.rgb = BADGE_COLORS.get(signal.classification, MID_GRAY)
    badge_run.font.name = FONT_BODY
    badge_run.bold = True
    _add_character_spacing(badge_run, 40)

    name_para.add_run("   ")

    score_run = name_para.add_run(f"{signal.overlap_score:.0%} overlap")
    score_run.font.size = Pt(8.5)
    score_run.font.color.rgb = MID_GRAY
    score_run.font.name = FONT_BODY

    _add_body_text(doc, signal.description, size=Pt(9), color=MID_GRAY)

    url_para = doc.add_paragraph()
    url_para.paragraph_format.space_after = Pt(4)
    run = url_para.add_run(signal.url)
    run.font.size = Pt(7.5)
    run.font.color.rgb = MID_GRAY
    run.font.name = FONT_BODY


def _add_metadata_section(
    doc: Document,
    queries_run: int,
    total_results: int,
    new_count: int,
    updated_count: int,
) -> None:
    """Add the scan metadata section."""
    _add_colored_divider(doc, NAVY)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(8)
    run = heading.add_run("SCAN DETAILS")
    run.font.size = Pt(8)
    run.font.color.rgb = NAVY
    run.font.name = FONT_BODY
    run.bold = True
    _add_character_spacing(run, 80)

    metadata_lines = [
        f"Scan Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
        f"Queries Run: {queries_run}",
        f"Total Results Searched: {total_results}",
        f"New Signals: {new_count}",
        f"Updated Signals: {updated_count}",
    ]

    for line in metadata_lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        run = para.add_run(line)
        run.font.size = Pt(8.5)
        run.font.color.rgb = MID_GRAY
        run.font.name = FONT_BODY


def _add_footer(doc: Document) -> None:
    """Add the report footer with wordmark."""
    _add_spacer(doc)
    _add_colored_divider(doc, NAVY)

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(8)

    run = footer_para.add_run("Generated by ")
    run.font.size = Pt(7.5)
    run.font.color.rgb = MID_GRAY
    run.font.name = FONT_BODY
    run.italic = True

    _add_wordmark(footer_para, Pt(7.5), MID_GRAY)

    run = footer_para.add_run(" Competitive Scanner")
    run.font.size = Pt(7.5)
    run.font.color.rgb = MID_GRAY
    run.font.name = FONT_BODY
    run.italic = True


# ═════════════════════════════════════════════════════════════════════════════
#  FORMATTING HELPERS
# ═════════════════════════════════════════════════════════════════════════════

def _add_body_text(
    doc: Document,
    text: str,
    size: Pt = Pt(10),
    color: RGBColor = DARK_GRAY,
) -> None:
    """Add a simple text paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.size = size
    run.font.color.rgb = color
    run.font.name = FONT_BODY


def _add_muted_text(doc: Document, text: str) -> None:
    """Add a muted description paragraph for section context."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = MID_GRAY
    run.font.name = FONT_BODY
    run.italic = True


def _add_spacer(doc: Document) -> None:
    """Add vertical space."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)


def _add_divider(doc: Document) -> None:
    """Add a light gray divider line between entries."""
    _add_colored_divider(doc, LIGHT_GRAY)


def _add_colored_divider(doc: Document, color: RGBColor) -> None:
    """Add a thin colored horizontal line."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): f"{color[0]:02X}{color[1]:02X}{color[2]:02X}",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_character_spacing(run, spacing_twips: int) -> None:
    """Add letter-spacing to a run for clean all-caps label look."""
    rPr = run._r.get_or_add_rPr()
    spacing_elem = rPr.makeelement(
        qn("w:spacing"),
        {qn("w:val"): str(spacing_twips)},
    )
    rPr.append(spacing_elem)
